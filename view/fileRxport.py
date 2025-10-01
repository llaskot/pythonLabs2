from datetime import datetime
from pathlib import Path
import openpyxl
from docx import Document


def export_to_xlsx(tree, filename="output"):
    Path("saved_files").mkdir(exist_ok=True)
    wb = openpyxl.Workbook()
    ws = wb.active
    cols = tree["columns"]
    ws.append(cols)
    for item_id in tree.get_children():
        row = tree.item(item_id)["values"]
        ws.append(row)
    now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    wb.save(f"saved_files/{filename}_{now}.xlsx")
    return True


def export_to_doc(tree, filename="output"):
    Path("saved_files").mkdir(exist_ok=True)
    doc = Document()
    doc.add_heading(filename, level=1)

    headers = tree["columns"]
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    for row_id in tree.get_children():
        row = tree.item(row_id)['values']
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    doc.save(f"saved_files/{filename}_{now}.docx")
    return True
